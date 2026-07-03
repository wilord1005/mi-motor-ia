import streamlit as st
import datetime
import json
import io
import re
import concurrent.futures
import requests
from docx import Document
from docx.shared import Pt
import google.generativeai as genai
from groq import Groq
from mistralai import Mistral
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
from google.oauth2 import service_account
from duckduckgo_search import DDGS
from bs4 import BeautifulSoup

# ============================================================
# CONFIG
# ============================================================
st.set_page_config(page_title="Motor Alfa Bravo", page_icon="✈️", layout="wide", initial_sidebar_state="collapsed")

FASES = {
    "F0": "00_Identidad_Motor.docx",
    "F1": "F1.docx", "F2": "F2.docx", "F3": "F3.docx",
    "F4": "F4.docx", "F4.1": "F4.1.docx", "F4.2": "F4.2.docx",
    "F4.3": "F4.3.docx", "F4.4": "F4.4.docx", "F4.5": "F4.5.docx",
    "F4.6": "F4.6.docx", "F5": "F5.docx", "F6": "F6.docx"
}
FASE_NOMBRES = {
    "F0":"Identidad","F1":"Recolección","F2":"Estadística","F3":"Comparación",
    "F4":"Simulación","F4.1":"Estrés","F4.2":"Escenarios","F4.3":"Marcadores",
    "F4.4":"Coherencia","F4.5":"Riesgo","F4.6":"Comparativa","F5":"Picks","F6":"Auditoría"
}

# ============================================================
# CSS
# ============================================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Share+Tech+Mono&family=Rajdhani:wght@400;600;700&display=swap');
html,body,[data-testid="stApp"]{background:#060d18;color:#c8d8e8;font-family:'Rajdhani',sans-serif;}
.hdr{background:linear-gradient(90deg,#070e1a,#0d1e35,#070e1a);border-bottom:2px solid #1e90ff;padding:12px 20px;}
.hdr-title{font-family:'Share Tech Mono',monospace;font-size:18px;color:#1e90ff;letter-spacing:4px;}
.hdr-sub{font-size:10px;color:#2a4a6a;letter-spacing:2px;margin-top:2px;}
.fuente-card{border-radius:4px;padding:8px 12px;margin:4px 0;font-size:12px;line-height:1.5;}
.f-python{background:#071207;border-left:3px solid #00ff88;color:#80d880;}
.f-odds{background:#120c07;border-left:3px solid #f4a300;color:#d4a060;}
.f-football{background:#070d12;border-left:3px solid #00bfff;color:#60b8d8;}
.f-sportsdb{background:#0a0712;border-left:3px solid #c060ff;color:#b080d8;}
.f-gemini{background:#07071a;border-left:3px solid #4285f4;color:#8090f0;}
.f-groq{background:#120707;border-left:3px solid #ff6b35;color:#e09070;}
.f-mistral{background:#0d0712;border-left:3px solid #9b59b6;color:#c090d8;}
.f-usuario{background:#07101a;border-left:3px solid #1e90ff;color:#c8d8e8;}
.debate-box{background:#080f1a;border:1px dashed #1a3a5c;border-radius:4px;padding:10px 14px;margin:6px 0;font-size:11px;color:#4a7a9b;font-family:'Share Tech Mono',monospace;}
.consenso-box{background:linear-gradient(135deg,#071207,#07101a);border:2px solid #1e90ff;border-radius:6px;padding:14px 16px;margin:10px 0;}
.consenso-titulo{font-family:'Share Tech Mono',monospace;font-size:11px;color:#1e90ff;letter-spacing:3px;margin-bottom:8px;}
.badge{display:inline-block;border-radius:2px;padding:1px 6px;font-size:10px;font-family:'Share Tech Mono',monospace;margin-right:4px;}
.b-python{background:#071207;color:#00ff88;border:1px solid #00ff88;}
.b-odds{background:#120c07;color:#f4a300;border:1px solid #f4a300;}
.b-football{background:#070d12;color:#00bfff;border:1px solid #00bfff;}
.b-sportsdb{background:#0a0712;color:#c060ff;border:1px solid #c060ff;}
.b-gemini{background:#07071a;color:#4285f4;border:1px solid #4285f4;}
.b-groq{background:#120707;color:#ff6b35;border:1px solid #ff6b35;}
.b-mistral{background:#0d0712;color:#9b59b6;border:1px solid #9b59b6;}
.b-usuario{background:#07101a;color:#1e90ff;border:1px solid #1e90ff;}
.b-consenso{background:#071a07;color:#00ff88;border:1px solid #00ff88;}
.alerta{background:#120707;border:1px solid #ff3333;border-radius:3px;padding:6px 10px;font-size:11px;color:#ff8888;font-family:'Share Tech Mono',monospace;}
.aprendizaje{background:#071207;border:1px dashed #00ff88;border-radius:3px;padding:6px 10px;font-size:11px;color:#60c870;}
.sdot{display:inline-block;width:7px;height:7px;border-radius:50%;margin-right:5px;animation:pulse 2s infinite;}
.don{background:#00ff88;} .doff{background:#333;}
@keyframes pulse{0%,100%{opacity:1;}50%{opacity:.3;}}
.tab-btn{background:#0a1422;border:1px solid #1a3a5c;border-radius:3px;padding:5px 10px;font-size:11px;cursor:pointer;color:#6a9acb;font-family:'Share Tech Mono',monospace;width:100%;text-align:left;margin:1px 0;}
.tab-active{border-color:#1e90ff!important;color:#1e90ff!important;}
.seccion-header{font-family:'Share Tech Mono',monospace;font-size:10px;color:#2a4a6a;letter-spacing:2px;padding:6px 0;border-bottom:1px solid #1a3a5c;margin-bottom:8px;}
.input-area{background:#060d18;border:1px solid #1a3a5c;border-radius:4px;padding:8px;}
.consola-central{background:linear-gradient(135deg,#060d18,#07101a);border:2px solid #1e90ff;border-radius:6px;}
.cc-header{background:#0a1a2a;border-bottom:1px solid #1e90ff;padding:10px 16px;border-radius:6px 6px 0 0;}
.cc-title{font-family:'Share Tech Mono',monospace;font-size:13px;color:#1e90ff;letter-spacing:3px;}
.cc-sub{font-size:10px;color:#2a4a6a;margin-top:2px;}
div[data-testid="stChatInput"] textarea{background:#060d18!important;color:#c8d8e8!important;font-family:'Share Tech Mono',monospace!important;border:1px solid #1a3a5c!important;font-size:13px!important;}
.stButton>button{background:#0a1422!important;border:1px solid #1a3a5c!important;color:#c8d8e8!important;font-family:'Rajdhani',sans-serif!important;border-radius:3px!important;}
.stButton>button:hover{border-color:#1e90ff!important;color:#1e90ff!important;}
hr{border-color:#1a3a5c!important;}
.stTabs [data-baseweb="tab-list"]{background:#060d18;border-bottom:1px solid #1a3a5c;}
.stTabs [data-baseweb="tab"]{background:#060d18;color:#4a7a9b;font-family:'Share Tech Mono',monospace;font-size:11px;}
.stTabs [aria-selected="true"]{background:#0a1422;color:#1e90ff;border-bottom:2px solid #1e90ff;}
</style>
""", unsafe_allow_html=True)

# ============================================================
# ESTADO
# ============================================================
def init():
    d = {
        "drive_service": None,
        "drive_folder_id": None,
        "fases_cache": {},
        "fase_activa": "F0",
        "datos_compartidos": {},
        "num_aprendizaje": 1,
        "picks_emitidos": [],
        "consola_mensajes": [],
        "fase_mensajes": {f: [] for f in FASES},
        "evento_activo": {},
    }
    for k, v in d.items():
        if k not in st.session_state:
            st.session_state[k] = v
init()

# ============================================================
# GOOGLE DRIVE
# ============================================================
def conectar_drive():
    try:
        creds = service_account.Credentials.from_service_account_info(
            dict(st.secrets["google_service_account"]),
            scopes=["https://www.googleapis.com/auth/drive"]
        )
        svc = build("drive", "v3", credentials=creds)
        st.session_state["drive_service"] = svc
        res = svc.files().list(
            q="name='MOTOR_PREDICCION' and mimeType='application/vnd.google-apps.folder' and trashed=false",
            fields="files(id)"
        ).execute()
        files = res.get("files", [])
        if files:
            st.session_state["drive_folder_id"] = files[0]["id"]
        return svc
    except Exception as e:
        st.session_state["drive_error"] = str(e)
        return None

def leer_docx(service, nombre):
    if nombre in st.session_state["fases_cache"]:
        return st.session_state["fases_cache"][nombre]
    try:
        folder_id = st.session_state.get("drive_folder_id")
        q = f"name='{nombre}' and trashed=false"
        if folder_id:
            q += f" and '{folder_id}' in parents"
        res = service.files().list(q=q, fields="files(id)").execute()
        files = res.get("files", [])
        if not files:
            return f"[{nombre} NO ENCONTRADO]"
        buf = io.BytesIO()
        dl = MediaIoBaseDownload(buf, service.files().get_media(fileId=files[0]["id"]))
        done = False
        while not done:
            _, done = dl.next_chunk()
        buf.seek(0)
        doc = Document(buf)
        texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        st.session_state["fases_cache"][nombre] = texto
        return texto
    except Exception as e:
        return f"[ERROR: {e}]"

def escribir_docx(service, nombre, contenido):
    try:
        folder_id = st.session_state.get("drive_folder_id")
        q = f"name='{nombre}' and trashed=false"
        if folder_id:
            q += f" and '{folder_id}' in parents"
        res = service.files().list(q=q, fields="files(id)").execute()
        files = res.get("files", [])
        if not files:
            return False
        file_id = files[0]["id"]
        buf = io.BytesIO()
        dl = MediaIoBaseDownload(buf, service.files().get_media(fileId=file_id))
        done = False
        while not done:
            _, done = dl.next_chunk()
        buf.seek(0)
        doc = Document(buf)
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        doc.add_paragraph(f"[{ts}] {contenido}")
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        service.files().update(
            fileId=file_id,
            media_body=MediaIoBaseUpload(out, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        ).execute()
        if nombre in st.session_state["fases_cache"]:
            del st.session_state["fases_cache"][nombre]
        return True
    except:
        return False

# ============================================================
# APIS DEPORTIVAS
# ============================================================
def api_odds(evento_query):
    """The Odds API — cuotas en tiempo real."""
    try:
        key = st.secrets.get("ODDS_API_KEY", "")
        if not key:
            return "[ODDS API: sin key]"
        url = "https://api.the-odds-api.com/v4/sports/soccer/odds"
        params = {"apiKey": key, "regions": "eu", "markets": "h2h", "oddsFormat": "decimal"}
        r = requests.get(url, params=params, timeout=8)
        data = r.json()
        if not isinstance(data, list):
            return f"[ODDS API: {data}]"
        texto = ""
        for ev in data[:5]:
            home = ev.get("home_team", "")
            away = ev.get("away_team", "")
            commence = ev.get("commence_time", "")
            if evento_query.lower() in home.lower() or evento_query.lower() in away.lower() or not evento_query:
                bookmakers = ev.get("bookmakers", [])
                cuotas = ""
                if bookmakers:
                    outcomes = bookmakers[0].get("markets", [{}])[0].get("outcomes", [])
                    cuotas = " | ".join(f"{o['name']}: {o['price']}" for o in outcomes)
                texto += f"• {home} vs {away} | {commence[:10]} | {cuotas}\n"
        return texto if texto else "[ODDS API: sin partidos para esa búsqueda]"
    except Exception as e:
        return f"[ODDS API ERROR: {e}]"

def api_football(endpoint, params_extra={}):
    """API-Football — estadísticas."""
    try:
        key = st.secrets.get("FOOTBALL_API_KEY", "")
        if not key:
            return "[FOOTBALL API: sin key]"
        headers = {"x-apisports-key": key}
        url = f"https://v3.football.api-sports.io/{endpoint}"
        r = requests.get(url, headers=headers, params=params_extra, timeout=8)
        data = r.json()
        resp = data.get("response", [])
        if not resp:
            return f"[FOOTBALL API: sin datos para {endpoint}]"
        texto = json.dumps(resp[:3], ensure_ascii=False, indent=2)
        return texto[:1500]
    except Exception as e:
        return f"[FOOTBALL API ERROR: {e}]"

def api_sportsdb(endpoint):
    """TheSportsDB — gratis sin key."""
    try:
        url = f"https://www.thesportsdb.com/api/v1/json/3/{endpoint}"
        r = requests.get(url, timeout=8)
        data = r.json()
        return json.dumps(list(data.values())[0][:3] if data else {}, ensure_ascii=False, indent=2)[:1000]
    except Exception as e:
        return f"[SPORTSDB ERROR: {e}]"

def python_scraping(query):
    """Python hace scraping independiente."""
    texto = ""
    try:
        with DDGS() as ddgs:
            raw = list(ddgs.text(query + " estadísticas fútbol resultados", max_results=6))
        for r in raw:
            texto += f"• {r.get('title','')}: {r.get('body','')}\n  [{r.get('href','')}]\n"
        if raw:
            try:
                resp = requests.get(raw[0]["href"], timeout=5, headers={"User-Agent": "Mozilla/5.0"})
                soup = BeautifulSoup(resp.text, "html.parser")
                tablas = soup.find_all("table")
                for t in tablas[:1]:
                    for fila in t.find_all("tr")[:6]:
                        celdas = fila.find_all(["td","th"])
                        if celdas:
                            texto += " | ".join(c.get_text(strip=True) for c in celdas) + "\n"
            except:
                pass
    except:
        pass
    return texto[:1500] if texto else "[Python: sin resultados]"

# ============================================================
# IAs
# ============================================================
def gemini(prompt):
    try:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        return genai.GenerativeModel("gemini-flash-latest").generate_content(prompt).text
    except Exception as e:
        return f"[GEMINI ERROR: {e}]"

def groq_ia(prompt):
    try:
        c = Groq(api_key=st.secrets["GROQ_API_KEY"])
        r = c.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role":"user","content":prompt}],
            max_tokens=3000, temperature=0.1
        )
        return r.choices[0].message.content
    except Exception as e:
        return f"[GROQ ERROR: {e}]"

def mistral_ia(prompt):
    try:
        c = Mistral(api_key=st.secrets["MISTRAL_API_KEY"])
        r = c.chat.complete(
            model="mistral-small-latest",
            messages=[{"role": "user", "content": prompt}]
        )
        return r.choices[0].message.content
    except Exception as e:
        return f"[MISTRAL ERROR: {e}]"

def ia_busca(query, fn_ia):
    try:
        with DDGS() as ddgs:
            raw = list(ddgs.text(query, max_results=5))
        return "\n".join(f"• {r.get('title','')}: {r.get('body','')}" for r in raw)
    except:
        return "[sin resultados]"

# ============================================================
# MOTOR CUADRANGULAR CON APIs
# ============================================================
def motor_completo(comando, fase, doc_motor, datos_evento=""):
    ts = datetime.datetime.now().strftime("%H:%M:%S")
    evento_q = datos_evento or comando

    resultado = {
        "python": "", "odds": "", "football": "", "sportsdb": "",
        "gemini": "", "groq": "", "mistral": "",
        "debate1": "", "debate2": "", "consenso": "",
        "aprendizajes": [], "alertas": []
    }

    def buscar_python():
        return python_scraping(evento_q)

    def buscar_odds():
        equipo = evento_q.split("vs")[0].strip() if "vs" in evento_q.lower() else evento_q[:20]
        return api_odds(equipo)

    def buscar_football():
        if "vs" in evento_q.lower():
            partes = evento_q.lower().split("vs")
            equipo = partes[0].strip()
            r = api_football("teams", {"search": equipo})
            return r
        return api_football("fixtures", {"date": datetime.date.today().isoformat(), "timezone": "America/Guayaquil"})

    def buscar_sportsdb():
        equipo = evento_q.split("vs")[0].strip() if "vs" in evento_q.lower() else evento_q[:20]
        return api_sportsdb(f"searchteams.php?t={equipo.replace(' ', '_')}")

    def buscar_gemini():
        return ia_busca(f"{evento_q} estadísticas forma reciente resultados", gemini)

    def buscar_groq():
        return ia_busca(f"{evento_q} head to head historial enfrentamientos", groq_ia)

    def buscar_mistral():
        return ia_busca(f"{evento_q} tabla posiciones liga temporada", mistral_ia)

    with concurrent.futures.ThreadPoolExecutor(max_workers=7) as ex:
        fp = ex.submit(buscar_python)
        fo = ex.submit(buscar_odds)
        ff = ex.submit(buscar_football)
        fs = ex.submit(buscar_sportsdb)
        fg = ex.submit(buscar_gemini)
        fgr = ex.submit(buscar_groq)
        fm = ex.submit(buscar_mistral)

        resultado["python"] = fp.result()
        resultado["odds"] = fo.result()
        resultado["football"] = ff.result()
        resultado["sportsdb"] = fs.result()
        resultado["gemini_web"] = fg.result()
        resultado["groq_web"] = fgr.result()
        resultado["mistral_web"] = fm.result()

    st.session_state["datos_compartidos"][fase] = {
        "evento": evento_q,
        "python": resultado["python"],
        "odds": resultado["odds"],
        "football": resultado["football"],
        "sportsdb": resultado["sportsdb"],
        "timestamp": ts
    }

    datos_previos = ""
    for f_prev, d in st.session_state["datos_compartidos"].items():
        if f_prev != fase:
            datos_previos += f"\n=== DATOS DE {f_prev} (fase anterior) ===\n{json.dumps(d, ensure_ascii=False)[:400]}\n"

    prompt_analisis = f"""Eres parte del equipo analista del Motor Alfa Bravo.

=== DOCUMENTO OFICIAL DEL MOTOR (LEY INAMOVIBLE) ===
{doc_motor}

=== DATOS RECOPILADOS POR TODAS LAS FUENTES ===
PYTHON (scraping web): {resultado["python"][:600]}
ODDS API (cuotas reales): {resultado["odds"][:400]}
FOOTBALL API (stats): {resultado["football"][:600]}
SPORTSDB (fixtures): {resultado["sportsdb"][:400]}

=== DATOS DE FASES ANTERIORES (interconexión) ===
{datos_previos[:600]}

=== EVENTO ===
{evento_q}

=== FASE A EJECUTAR ===
{fase} — {FASE_NOMBRES.get(fase,"")}

=== TUS DATOS DE BÚSQUEDA WEB PROPIOS ===
{{DATOS_PROPIOS}}

INSTRUCCIONES:
- Aplica EXACTAMENTE las reglas del documento Word del motor.
- Usa SOLO datos con fuente verificada. Sin fuente = DATO AUSENTE.
- Cita la regla específica: "Según Regla F2-02..."
- Si encuentras algo que el motor no contempla: [LAGUNA DETECTADA: descripción]
- Sé específico. No repitas el documento — aplícalo al evento real.
- Eres {{ROL}}."""

    def analisis_gemini():
        p = prompt_analisis.replace("{DATOS_PROPIOS}", resultado["gemini_web"]).replace("{ROL}", "GEMINI — Analista 1")
        return gemini(p)

    def analisis_groq():
        p = prompt_analisis.replace("{DATOS_PROPIOS}", resultado["groq_web"]).replace("{ROL}", "GROQ — Analista 2")
        return groq_ia(p)

    def analisis_mistral():
        p = prompt_analisis.replace("{DATOS_PROPIOS}", resultado["mistral_web"]).replace("{ROL}", "MISTRAL — Analista 3")
        return mistral_ia(p)

    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as ex:
        ag = ex.submit(analisis_gemini)
        agr = ex.submit(analisis_groq)
        am = ex.submit(analisis_mistral)
        resultado["gemini"] = ag.result()
        resultado["groq"] = agr.result()
        resultado["mistral"] = am.result()

    prompt_d1 = f"""MOTOR ALFA BRAVO — DEBATE RONDA 1 — {fase}

DOCUMENTO DEL MOTOR (árbitro metodológico):
{doc_motor[:1500]}

DATOS DE TODAS LAS FUENTES:
Python: {resultado["python"][:400]}
Odds API: {resultado["odds"][:300]}
Football API: {resultado["football"][:400]}
SportsDB: {resultado["sportsdb"][:300]}

ANÁLISIS DE LOS TRES ANALISTAS:
GEMINI: {resultado["gemini"][:600]}
GROQ: {resultado["groq"][:600]}
MISTRAL: {resultado["mistral"][:600]}

Eres {{ROL}}.
DEBATE CON ARGUMENTOS Y DATOS REALES:
1. ¿Qué datos encontraron todos y coinciden? → alta confianza
2. ¿Qué datos difieren entre fuentes? → debate con fuente
3. ¿Qué aportó cada analista que los otros no vieron?
4. ¿Hay lagunas en el motor Word? → señalar con [LAGUNA]
5. ¿Qué corriges de tu propio análisis tras leer a los demás?
Habla directamente a los otros por su nombre. Argumenta con datos reales."""

    def d1_gemini():
        return gemini(prompt_d1.replace("{ROL}", "GEMINI"))
    def d1_groq():
        return groq_ia(prompt_d1.replace("{ROL}", "GROQ"))
    def d1_mistral():
        return mistral_ia(prompt_d1.replace("{ROL}", "MISTRAL"))

    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as ex:
        r1g = ex.submit(d1_gemini)
        r1gr = ex.submit(d1_groq)
        r1m = ex.submit(d1_mistral)
        d1g = r1g.result()
        d1gr = r1gr.result()
        d1m = r1m.result()
    resultado["debate1"] = f"GEMINI:\n{d1g}\n\nGROQ:\n{d1gr}\n\nMISTRAL:\n{d1m}"

    prompt_d2 = f"""MOTOR ALFA BRAVO — DEBATE RONDA 2 — CONSTRUIR CONSENSO — {fase}

RONDA 1 COMPLETA:
{resultado["debate1"][:2000]}

DATOS PYTHON + APIs:
{resultado["python"][:300]} | {resultado["odds"][:200]} | {resultado["football"][:300]}

DOCUMENTO DEL MOTOR:
{doc_motor[:1000]}

Eres {{ROL}}.
RONDA 2 — CONSTRUCCIÓN DE CONSENSO:
1. Puntos donde los 4 (Python+Gemini+Groq+Mistral) estamos de acuerdo → consenso sólido
2. Discrepancias que quedan → ¿cuál fuente es más confiable y por qué?
3. Lagunas detectadas → propuesta concreta para incorporar al motor
4. Tu posición final para el consenso de {fase}
El Word del motor es árbitro metodológico.
Los datos web verificados son árbitro de hechos.
Solo ceder cuando el argumento o dato lo justifica — no por cortesía."""

    def d2_gemini():
        return gemini(prompt_d2.replace("{ROL}", "GEMINI"))
    def d2_groq():
        return groq_ia(prompt_d2.replace("{ROL}", "GROQ"))
    def d2_mistral():
        return mistral_ia(prompt_d2.replace("{ROL}", "MISTRAL"))

    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as ex:
        r2g = ex.submit(d2_gemini)
        r2gr = ex.submit(d2_groq)
        r2m = ex.submit(d2_mistral)
        d2g = r2g.result()
        d2gr = r2gr.result()
        d2m = r2m.result()
    resultado["debate2"] = f"GEMINI:\n{d2g}\n\nGROQ:\n{d2gr}\n\nMISTRAL:\n{d2m}"

    prompt_consenso = f"""Eres Python — sistema nervioso y supervisor del Motor Alfa Bravo.

FASE: {fase} — {FASE_NOMBRES.get(fase,"")}
EVENTO: {evento_q}

ANÁLISIS INDIVIDUALES:
Gemini: {resultado["gemini"][:500]}
Groq: {resultado["groq"][:500]}
Mistral: {resultado["mistral"][:500]}

DEBATE COMPLETO:
Ronda 1: {resultado["debate1"][:800]}
Ronda 2: {resultado["debate2"][:800]}

DATOS VERIFICADOS:
Python: {resultado["python"][:400]}
Odds API: {resultado["odds"][:300]}
Football API: {resultado["football"][:300]}

DOCUMENTO DEL MOTOR:
{doc_motor[:1200]}

TU ROL COMO PYTHON-SUPERVISOR:
1. VALIDAR que el consenso respeta el Word del motor.
2. ESTRUCTURAR el producto oficial de {fase}.
3. MARCAR datos sin verificar como ALERTA.
4. DETECTAR lagunas y proponer aprendizajes F6.
5. INTERCONECTAR: qué datos de {fase} alimentan la siguiente fase.

FORMATO OBLIGATORIO:
--- PRODUCTO OFICIAL {fase} ---
[resultado verificado por los 4 participantes]

--- DATOS PARA SIGUIENTE FASE ---
[qué información pasa a la siguiente fase]

--- ALERTAS ---
[datos discrepantes o sin verificar]

--- LAGUNAS DETECTADAS ---
[lo que el motor Word no contempla]

--- APRENDIZAJES PROPUESTOS ---
Aprendizaje F6-XXX | Fase afectada: FX | Regla afectada: FX-XX | Corrección: descripción

Solo hechos verificados. Nada inventado."""

    consenso = gemini(prompt_consenso)
    if "ERROR" in consenso:
        consenso = groq_ia(prompt_consenso)
    resultado["consenso"] = consenso

    st.session_state["datos_compartidos"][fase]["consenso"] = consenso[:800]

    for linea in consenso.split("\n"):
        if "Aprendizaje F6-" in linea:
            resultado["aprendizajes"].append(linea.strip())
    en_alertas = False
    for linea in consenso.split("\n"):
        if "ALERTAS" in linea:
            en_alertas = True
        elif "---" in linea:
            en_alertas = False
        elif en_alertas and linea.strip():
            resultado["alertas"].append(linea.strip())

    return resultado

# ============================================================
# CARGAR DOCUMENTOS MOTOR
# ============================================================
def cargar_docs_fase(service, fase):
    if not service:
        return "Drive no conectado."
    fases_a_cargar = ["F0", fase]
    if fase not in ["F0", "F1"]:
        fases_a_cargar.extend(["F1"])
    if fase in ["F3","F4","F4.1","F4.2","F4.3","F4.4","F4.5","F4.6"]:
        fases_a_cargar.extend(["F2","F3"])
    if fase in ["F5","F6"]:
        fases_a_cargar.extend(["F2","F3","F4"])
    docs = {}
    for fk in list(dict.fromkeys(fases_a_cargar)):
        if fk in FASES:
            docs[fk] = leer_docx(service, FASES[fk])
    return "\n\n".join(f"=== {fk} — {FASE_NOMBRES.get(fk,'')} ===\n{txt}" for fk, txt in docs.items())

# ============================================================
# RENDERIZAR RESULTADO
# ============================================================
def render_resultado(r, fase, hora):
    st.markdown(f'<div class="seccion-header">— ANÁLISIS {fase} · {hora} —</div>', unsafe_allow_html=True)

    cols = st.columns(4)
    with cols[0]:
        st.markdown(f'<div class="fuente-card f-python"><span class="badge b-python">PYTHON</span><br><small>{r["python"][:200]}...</small></div>', unsafe_allow_html=True)
    with cols[1]:
        st.markdown(f'<div class="fuente-card f-odds"><span class="badge b-odds">ODDS API</span><br><small>{r["odds"][:200]}...</small></div>', unsafe_allow_html=True)
    with cols[2]:
        st.markdown(f'<div class="fuente-card f-football"><span class="badge b-football">FOOTBALL API</span><br><small>{r["football"][:200]}...</small></div>', unsafe_allow_html=True)
    with cols[3]:
        st.markdown(f'<div class="fuente-card f-sportsdb"><span class="badge b-sportsdb">SPORTSDB</span><br><small>{r["sportsdb"][:200]}...</small></div>', unsafe_allow_html=True)

    cols2 = st.columns(3)
    with cols2[0]:
        st.markdown(f'<div class="fuente-card f-gemini"><span class="badge b-gemini">GEMINI</span><br>{r["gemini"][:500]}</div>', unsafe_allow_html=True)
    with cols2[1]:
        st.markdown(f'<div class="fuente-card f-groq"><span class="badge b-groq">GROQ</span><br>{r["groq"][:500]}</div>', unsafe_allow_html=True)
    with cols2[2]:
        st.markdown(f'<div class="fuente-card f-mistral"><span class="badge b-mistral">MISTRAL</span><br>{r["mistral"][:500]}</div>', unsafe_allow_html=True)

    with st.expander("💬 Debate completo entre los 4 participantes"):
        st.markdown(f'<div class="debate-box"><b style="color:#1e90ff">RONDA 1</b><br>{r["debate1"][:1500]}<br><br><b style="color:#1e90ff">RONDA 2</b><br>{r["debate2"][:1500]}</div>', unsafe_allow_html=True)

    st.markdown(f'<div class="consenso-box"><div class="consenso-titulo">◆ CONSENSO CUADRANGULAR — {fase}</div><div>{r["consenso"]}</div></div>', unsafe_allow_html=True)

    for a in r.get("alertas", []):
        st.markdown(f'<div class="alerta">⚠ {a}</div>', unsafe_allow_html=True)

    for ap in r.get("aprendizajes", []):
        st.markdown(f'<div class="aprendizaje">📚 {ap}</div>', unsafe_allow_html=True)

# ============================================================
# INTERFAZ PRINCIPAL
# ============================================================
st.markdown("""
<div class="hdr">
    <div class="hdr-title">✈ MOTOR ALFA BRAVO</div>
    <div class="hdr-sub">Python · Gemini · Groq · Mistral · Odds API · Football API · TheSportsDB · Drive</div>
</div>
""", unsafe_allow_html=True)

c1,c2,c3,c4,c5 = st.columns([2,2,2,2,2])
drive_ok = st.session_state.get("drive_service") is not None
with c1:
    d = "don" if drive_ok else "doff"
    st.markdown(f'<div style="font-family:\'Share Tech Mono\',monospace;font-size:10px;padding:8px 0;"><span class="sdot {d}"></span>{"Drive conectado" if drive_ok else "Drive off"}</div>', unsafe_allow_html=True)
with c2:
    fa = st.session_state.get("fase_activa","F0")
    st.markdown(f'<div style="font-family:\'Share Tech Mono\',monospace;font-size:10px;padding:8px 0;color:#4a7a9b;">FASE: <span style="color:#1e90ff">{fa}</span></div>', unsafe_allow_html=True)
with c3:
    picks = len(st.session_state.get("picks_emitidos",[]))
    st.markdown(f'<div style="font-family:\'Share Tech Mono\',monospace;font-size:10px;padding:8px 0;color:#4a7a9b;">PICKS: <span style="color:#00ff88">{picks}</span></div>', unsafe_allow_html=True)
with c4:
    fases_cargadas = len(st.session_state.get("fases_cache",{}))
    st.markdown(f'<div style="font-family:\'Share Tech Mono\',monospace;font-size:10px;padding:8px 0;color:#4a7a9b;">DOCS: <span style="color:#f4a300">{fases_cargadas}</span></div>', unsafe_allow_html=True)
with c5:
    if not drive_ok:
        if st.button("🔗 Conectar Drive", use_container_width=True):
            with st.spinner("Conectando..."):
                svc = conectar_drive()
                if svc:
                    st.session_state["consola_mensajes"].append({
                        "rol":"sistema","texto":"✅ Google Drive conectado. Motor Alfa Bravo en línea.","hora":datetime.datetime.now().strftime("%H:%M")
                    })
                    st.rerun()
                else:
                    st.error(f"Error: {st.session_state.get('drive_error', 'desconocido')}")
    else:
        if st.button("🔄 Recargar docs", use_container_width=True):
            st.session_state["fases_cache"] = {}
            st.rerun()

st.markdown("<hr>", unsafe_allow_html=True)

nombres_tabs = ["🎛 CONSOLA CENTRAL"] + [f"{fk} · {FASE_NOMBRES[fk]}" for fk in FASES.keys()]
tabs = st.tabs(nombres_tabs)

with tabs[0]:
    st.markdown("""
<div class="consola-box">
<div class="cc-header">
    <div class="cc-title">🎛 CONSOLA CENTRAL — CONTROL TOTAL DEL MOTOR</div>
    <div class="cc-sub">Tú · Python · Gemini · Groq · Mistral — Equipo de 5 · Control absoluto</div>
</div>
</div>
""", unsafe_allow_html=True)

    st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)

    bc1,bc2,bc3,bc4,bc5 = st.columns(5)
    with bc1:
        if st.button("📋 Ver picks F5", use_container_width=True):
            picks = st.session_state.get("picks_emitidos",[])
            msg = f"Picks emitidos: {len(picks)}\n" + "\n".join(f"• {p['hora']} {p['fase']}: {p['resumen']}" for p in picks[-5:])
            st.session_state["consola_mensajes"].append({"rol":"sistema","texto":msg,"hora":datetime.datetime.now().strftime("%H:%M")})
            st.rerun()
    with bc2:
        if st.button("🔄 Limpiar sesión", use_container_width=True):
            st.session_state["datos_compartidos"] = {}
            st.session_state["evento_activo"] = {}
            st.session_state["consola_mensajes"].append({"rol":"sistema","texto":"Sesión limpiada. Datos compartidos borrados.","hora":datetime.datetime.now().strftime("%H:%M")})
            st.rerun()
    with bc3:
        if st.button("📊 Estado fases", use_container_width=True):
            dc = st.session_state.get("datos_compartidos",{})
            estado = "\n".join(f"• {fk}: {'✅ analizada' if fk in dc else '⏳ pendiente'}" for fk in FASES.keys())
            st.session_state["consola_mensajes"].append({"rol":"sistema","texto":f"Estado:\n{estado}","hora":datetime.datetime.now().strftime("%H:%M")})
            st.rerun()
    with bc4:
        if st.button("📚 Ver aprendizajes", use_container_width=True):
            n = st.session_state.get("num_aprendizaje",1)
            st.session_state["consola_mensajes"].append({"rol":"sistema","texto":f"Aprendizajes registrados en esta sesión: {n-1}","hora":datetime.datetime.now().strftime("%H:%M")})
            st.rerun()
    with bc5:
        if st.button("⚡ Análisis F0→F5", use_container_width=True):
            st.session_state["consola_mensajes"].append({"rol":"sistema","texto":"Escribe el partido en el chat y ejecutaré el ciclo completo F0→F5.","hora":datetime.datetime.now().strftime("%H:%M")})
            st.rerun()

    contenedor_cc = st.container(height=480)
    with contenedor_cc:
        if not st.session_state["consola_mensajes"]:
            st.markdown("""
<div style="text-align:center;padding:80px 20px;">
<div style="font-family:'Share Tech Mono',monospace;font-size:40px;color:#1a3a5c;">✈</div>
<div style="font-family:'Share Tech Mono',monospace;font-size:14px;color:#1e90ff;letter-spacing:4px;margin:10px 0;">MOTOR ALFA BRAVO</div>
<div style="font-size:12px;color:#2a4a6a;line-height:2.2;">
Consola Central · Control Total<br>
Somos 5: Tú + Python + Gemini + Groq + Mistral<br><br>
<span style="color:#1a3a5c">Comandos:</span><br>
· "Analiza Cork vs Treaty cuotas 1.28/5.33/8.00"<br>
· "Ejecuta F1" · "Muestra picks" · "Modifica F4.2-05"<br>
· "Busca partidos de hoy" · "Guarda aprendizaje en F6"<br>
· "¿Cuánto riesgo tiene el pick de St Pat?"
</div>
</div>""", unsafe_allow_html=True)

        for msg in st.session_state["consola_mensajes"]:
            rol = msg.get("rol","")
            hora = msg.get("hora","")
            texto = msg.get("texto","")
            resultado = msg.get("resultado")

            if rol == "usuario":
                st.markdown(f'<div class="fuente-card f-usuario"><span class="badge b-usuario">TÚ</span> {hora}<br><div style="margin-top:6px;">{texto}</div></div>', unsafe_allow_html=True)
            elif rol == "sistema":
                st.markdown(f'<div style="font-family:\'Share Tech Mono\',monospace;font-size:11px;color:#00ff88;padding:4px 12px;">⚡ {texto}</div>', unsafe_allow_html=True)
            elif rol == "resultado" and resultado:
                render_resultado(resultado, msg.get("fase",""), hora)

    cmd_cc = st.chat_input("Escribe aquí — Tú + Python + Gemini + Groq + Mistral trabajamos juntos...", key="cc_input")
    if cmd_cc:
        hora = datetime.datetime.now().strftime("%H:%M")
        st.session_state["consola_mensajes"].append({"rol":"usuario","texto":cmd_cc,"hora":hora})

        service = st.session_state.get("drive_service")
        cmd_low = cmd_cc.lower()

        fase_cmd = None
        for fk in FASES.keys():
            if fk.lower() in cmd_low:
                fase_cmd = fk
                break
        if not fase_cmd:
            if any(x in cmd_low for x in ["analiza","ejecuta","corre","vs","partido","cuota"]):
                fase_cmd = "F0"

        if fase_cmd:
            doc_motor = cargar_docs_fase(service, fase_cmd)
            datos_ev = ""
            ctx = st.session_state.get("evento_activo",{})
            if ctx:
                datos_ev = ctx.get("input","")
            if any(x in cmd_low for x in ["vs","cuota","local:","visitante:"]):
                st.session_state["evento_activo"] = {"input": cmd_cc}
                datos_ev = cmd_cc

            with st.spinner("Los 6 analizando simultáneamente..."):
                resultado = motor_completo(cmd_cc, fase_cmd, doc_motor, datos_ev)

            st.session_state["consola_mensajes"].append({
                "rol":"resultado","resultado":resultado,"fase":fase_cmd,"hora":hora
            })
            st.session_state["fase_activa"] = fase_cmd

            if resultado.get("aprendizajes") and service:
                for ap in resultado["aprendizajes"]:
                    n = st.session_state["num_aprendizaje"]
                    ap_f = ap.replace("XXX", str(n).zfill(3))
                    escribir_docx(service, FASES["F6"], ap_f)
                    st.session_state["num_aprendizaje"] += 1

            if "operativo" in resultado.get("consenso","").lower() and "%" in resultado.get("consenso",""):
                st.session_state["picks_emitidos"].append({
                    "hora":hora,"fase":fase_cmd,"resumen":resultado["consenso"][:120]
                })
        else:
            doc_id = leer_docx(service, FASES["F0"]) if service else "Drive no conectado."
            prompt_gral = f"""Equipo Motor Alfa Bravo respondiendo consulta.
Documento base: {doc_id[:800]}
Datos compartidos de sesión: {json.dumps(list(st.session_state["datos_compartidos"].keys()))}
Consulta: {cmd_cc}
Responde como equipo (Python + Gemini + Groq + Mistral).
Si es un comando al motor, ejecútalo. Si es una consulta, respóndela con datos reales."""
            with st.spinner("Consultando al equipo..."):
                with concurrent.futures.ThreadPoolExecutor(max_workers=3) as ex:
                    rg = ex.submit(gemini, prompt_gral)
                    rgr = ex.submit(groq_ia, prompt_gral)
                    rm = ex.submit(mistral_ia, prompt_gral)
                    rg_r = rg.result()
                    rgr_r = rgr.result()
                    rm_r = rm.result()
            resp_final = f"**GEMINI:** {rg_r[:400]}\n\n**GROQ:** {rgr_r[:400]}\n\n**MISTRAL:** {rm_r[:400]}"
            st.session_state["consola_mensajes"].append({"rol":"sistema","texto":resp_final,"hora":hora})
        st.rerun()

for idx, (fase_key, archivo) in enumerate(FASES.items()):
    with tabs[idx + 1]:
        st.markdown(f'<div class="seccion-header">{fase_key} — {FASE_NOMBRES[fase_key].upper()} · Todos buscan · Todos analizan · Todos debaten</div>', unsafe_allow_html=True)

        dc = st.session_state.get("datos_compartidos", {})
        if dc:
            fases_con_datos = [f for f in dc.keys() if f != fase_key]
            if fases_con_datos:
                st.markdown(f'<div style="font-size:10px;color:#2a4a6a;font-family:\'Share Tech Mono\',monospace;padding:4px 0;">🔗 Datos disponibles de: {" · ".join(fases_con_datos)}</div>', unsafe_allow_html=True)

        mensajes_fase = st.session_state["fase_mensajes"].get(fase_key, [])
        contenedor_fase = st.container(height=420)
        with contenedor_fase:
            if not mensajes_fase:
                st.markdown(f"""
<div style="text-align:center;padding:40px 20px;color:#1a3a5c;">
    <div style="font-family:'Share Tech Mono',monospace;font-size:13px;color:#1e90ff;">{fase_key} — {FASE_NOMBRES[fase_key]}</div>
    <div style="font-size:11px;margin-top:8px;color:#2a4a6a;">
        Escribe abajo para iniciar el análisis.<br>
        Python + Odds API + Football API + Gemini + Groq + Mistral buscarán simultáneamente.
    </div>
</div>""", unsafe_allow_html=True)
            for msg in mensajes_fase:
                if msg.get("rol") == "usuario":
                    st.markdown(f'<div class="fuente-card f-usuario"><span class="badge b-usuario">TÚ</span> {msg.get("hora","")}<br>{msg.get("texto","")}</div>', unsafe_allow_html=True)
                elif msg.get("resultado"):
                    render_resultado(msg["resultado"], fase_key, msg.get("hora",""))

        cmd_fase = st.chat_input(
            f"Escribe en {fase_key} — {FASE_NOMBRES[fase_key]}...",
            key=f"input_{fase_key}"
        )
        if cmd_fase:
            hora = datetime.datetime.now().strftime("%H:%M")
            service = st.session_state.get("drive_service")

            if fase_key not in st.session_state["fase_mensajes"]:
                st.session_state["fase_mensajes"][fase_key] = []

            st.session_state["fase_mensajes"][fase_key].append({
                "rol":"usuario","texto":cmd_fase,"hora":hora
            })

            if any(x in cmd_fase.lower() for x in ["vs","cuota","analiza","ejecuta","partido","busca","estadísticas"]):
                doc_motor = cargar_docs_fase(service, fase_key)
                datos_ev = cmd_fase
                if "vs" in cmd_fase.lower() or "cuota" in cmd_fase.lower():
                    st.session_state["evento_activo"] = {"input": cmd_fase}
                elif st.session_state.get("evento_activo"):
                    datos_ev = st.session_state["evento_activo"].get("input","") + " " + cmd_fase

                with st.spinner(f"Los 6 analizando {fase_key}..."):
                    resultado = motor_completo(cmd_fase, fase_key, doc_motor, datos_ev)

                st.session_state["fase_mensajes"][fase_key].append({
                    "rol":"resultado","resultado":resultado,"hora":hora
                })
                st.session_state["fase_activa"] = fase_key

                if resultado.get("aprendizajes") and service:
                    for ap in resultado["aprendizajes"]:
                        n = st.session_state["num_aprendizaje"]
                        escribir_docx(service, FASES["F6"], ap.replace("XXX", str(n).zfill(3)))
                        st.session_state["num_aprendizaje"] += 1
            else:
                doc_fase = leer_docx(service, archivo) if service else "Drive no conectado."
                dc_actual = st.session_state["datos_compartidos"].get(fase_key, {})
                prompt_q = f"""Motor Alfa Bravo — {fase_key} — {FASE_NOMBRES[fase_key]}

Documento de {fase_key}: {doc_fase[:800]}
Datos de sesión de {fase_key}: {json.dumps(dc_actual)[:400]}
Datos compartidos: {json.dumps({k:v.get("consenso","")[:200] for k,v in st.session_state["datos_compartidos"].items()})[:600]}

Consulta del usuario: {cmd_fase}
Responde aplicando las reglas del documento Word. Cita las reglas específicas."""
                with st.spinner("Consultando..."):
                    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as ex:
                        rg = ex.submit(gemini, prompt_q)
                        rgr = ex.submit(groq_ia, prompt_q)
                        rm = ex.submit(mistral_ia, prompt_q)
                        resp = f"**GEMINI:** {rg.result()[:300]}\n\n**GROQ:** {rgr.result()[:300]}\n\n**MISTRAL:** {rm.result()[:300]}"

                st.session_state["fase_mensajes"][fase_key].append({
                    "rol":"sistema","texto":resp,"hora":hora
                })
            st.rerun()
